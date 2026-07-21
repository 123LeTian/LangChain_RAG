from pathlib import Path

import pytest

from src.ingestion import Indexer
from src.ingestion.exceptions import (
    DocumentLoadError,
    EmbeddingValidationError,
    EmptyContentError,
    IndexingOperationError,
    UnsupportedDocumentTypeError,
)
from src.ingestion.splitter import split_document
from src.knowledge.exceptions import DuplicateDocumentError
from src.knowledge.repositories import (
    InMemoryChunkRepository,
    InMemoryDocumentRepository,
    InMemoryKnowledgeRepository,
)
from src.knowledge.service import KnowledgeService
from src.models.knowledge import DocumentStatus, KnowledgeBaseStatus
from src.retrieval.embeddings import BaseEmbedder, HashEmbedder
from src.retrieval.vector_index import InMemoryVectorIndex


class WrongCountEmbedder(BaseEmbedder):
    @property
    def model_name(self) -> str:
        return "wrong-count"

    @property
    def dimension(self) -> int:
        return 3

    def embed_texts(self, texts):
        return []


class WrongDimensionEmbedder(BaseEmbedder):
    @property
    def model_name(self) -> str:
        return "wrong-dimension"

    @property
    def dimension(self) -> int:
        return 3

    def embed_texts(self, texts):
        return [[1.0, 2.0] for _ in texts]


class FailingVectorIndex(InMemoryVectorIndex):
    def upsert(self, chunks, vectors):
        super().upsert(chunks, vectors)
        raise RuntimeError("simulated vector write failure")


def deterministic_splitter(document):
    return split_document(
        document,
        chunk_size=80,
        chunk_overlap=0,
        separators=["\n\n", "\n", " "],
    )


def make_pipeline(*, embedder=None, vector_index=None, kb_ids=("kb-1",)):
    knowledge_repository = InMemoryKnowledgeRepository()
    document_repository = InMemoryDocumentRepository()
    chunk_repository = InMemoryChunkRepository()
    service = KnowledgeService(
        knowledge_repository,
        document_repository,
        chunk_repository,
    )
    for kb_id in kb_ids:
        service.create_knowledge_base(
            owner_id="owner",
            name=kb_id,
            embedding_model="hash",
            kb_id=kb_id,
        )
    chosen_embedder = embedder or HashEmbedder(dim=16)
    chosen_index = vector_index or InMemoryVectorIndex(chosen_embedder)
    indexer = Indexer(
        knowledge_service=service,
        splitter=deterministic_splitter,
        embedder=chosen_embedder,
        vector_index=chosen_index,
    )
    return indexer, service, chosen_index


def write_minimal_pdf(path: Path, text: str = "Offline PDF text") -> None:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 14 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{number} 0 obj\n".encode("ascii"))
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(payload)


def assert_indexed(service, vector_index, document_id, kb_id="kb-1"):
    document = service.get_document(document_id)
    chunks = service.get_document_chunks(document_id)
    assert document.status == DocumentStatus.INDEXED
    assert service.get_knowledge_base(kb_id).status == KnowledgeBaseStatus.READY
    assert chunks
    assert vector_index.count() == len(chunks)
    for chunk in chunks:
        assert chunk.metadata["kb_id"] == kb_id
        assert chunk.metadata["document_id"] == document_id
        assert chunk.metadata["chunk_id"] == chunk.id
        assert chunk.metadata["filename"] == document.filename
        assert "page" in chunk.metadata
        assert "section" in chunk.metadata
        assert chunk.metadata["chunk_index"] == chunk.index
    query = vector_index.embedder.embed_query(document.text)
    hits = vector_index.search(
        query,
        kb_id=document.kb_id,
        top_k=len(chunks),
    )
    assert {hit.chunk.id for hit in hits} == {chunk.id for chunk in chunks}
    assert all(hit.chunk.metadata["chunk_id"] == hit.chunk.id for hit in hits)
    return document, chunks


def test_real_txt_runs_loader_normalizer_splitter_chunks_and_vectors(tmp_path):
    source = tmp_path / "notes.txt"
    source.write_bytes("\ufeff第一段\r\n\r\n第二   段\u200b".encode("utf-8"))
    indexer, service, vector_index = make_pipeline()

    result = indexer.index_document("kb-1", source, document_id="doc-txt")
    document, chunks = assert_indexed(service, vector_index, result.id)

    assert document.text == "第一段\n\n第二 段"
    assert "第一段" in "".join(chunk.text for chunk in chunks)
    assert chunks[0].metadata["source"] == "notes.txt"


def test_indexer_applies_the_complete_document_and_knowledge_status_flow(
    tmp_path, monkeypatch
):
    source = tmp_path / "states.txt"
    source.write_text("state transition content", encoding="utf-8")
    indexer, service, _ = make_pipeline()
    document_states = []
    knowledge_states = []

    original_document_update = service.document_repository.update
    original_document_status = service.document_repository.update_status
    original_knowledge_status = service.knowledge_repository.update_status

    def record_document_update(document):
        result = original_document_update(document)
        document_states.append(result.status)
        return result

    def record_document_status(document_id, status, error_message=None):
        result = original_document_status(document_id, status, error_message)
        document_states.append(result.status)
        return result

    def record_knowledge_status(kb_id, status, error_message=None):
        result = original_knowledge_status(kb_id, status, error_message)
        knowledge_states.append(result.status)
        return result

    monkeypatch.setattr(service.document_repository, "update", record_document_update)
    monkeypatch.setattr(
        service.document_repository, "update_status", record_document_status
    )
    monkeypatch.setattr(
        service.knowledge_repository, "update_status", record_knowledge_status
    )

    indexer.index_document("kb-1", source, document_id="doc-states")

    assert document_states == [
        DocumentStatus.PARSING,
        DocumentStatus.PARSED,
        DocumentStatus.INDEXING,
        DocumentStatus.INDEXED,
    ]
    assert knowledge_states == [
        KnowledgeBaseStatus.INDEXING,
        KnowledgeBaseStatus.READY,
    ]


def test_real_markdown_preserves_heading_paragraphs_and_section(tmp_path):
    source = tmp_path / "guide.md"
    source.write_text("# 安装\n\n第一段。\n\n第二段。", encoding="utf-8")
    indexer, service, vector_index = make_pipeline()

    result = indexer.index_document("kb-1", source, document_id="doc-md")
    document, chunks = assert_indexed(service, vector_index, result.id)

    assert document.text.startswith("# 安装\n\n")
    assert chunks[0].metadata["section"] == "安装"
    assert chunks[0].metadata["source"] == "guide.md"


def test_real_pdf_is_generated_parsed_and_keeps_page_metadata(tmp_path):
    source = tmp_path / "offline.pdf"
    write_minimal_pdf(source)
    indexer, service, vector_index = make_pipeline()

    result = indexer.index_document("kb-1", source, document_id="doc-pdf")
    document, chunks = assert_indexed(service, vector_index, result.id)

    assert "Offline PDF text" in document.text
    assert chunks[0].metadata["page"] == 1
    assert document.metadata["page_count"] == 1


def test_real_docx_is_generated_parsed_and_keeps_heading_metadata(tmp_path):
    from docx import Document

    source = tmp_path / "manual.docx"
    docx = Document()
    docx.add_heading("Runtime Guide", level=1)
    docx.add_paragraph("First paragraph")
    docx.add_paragraph("Second paragraph")
    docx.save(source)
    indexer, service, vector_index = make_pipeline()

    result = indexer.index_document("kb-1", source, document_id="doc-docx")
    document, chunks = assert_indexed(service, vector_index, result.id)

    assert "Runtime Guide" in document.text
    assert "First paragraph" in document.text
    assert "Second paragraph" in document.text
    assert chunks[0].metadata["section"] == "Runtime Guide"
    assert chunks[0].metadata["source"] == "manual.docx"


def test_checksum_uses_raw_bytes_and_deduplicates_only_within_kb(tmp_path):
    source = tmp_path / "same.txt"
    source.write_bytes(b"same raw bytes")
    indexer, service, vector_index = make_pipeline(kb_ids=("kb-1", "kb-2"))

    first = indexer.index_document("kb-1", source, document_id="doc-1")
    count = vector_index.count()
    with pytest.raises(DuplicateDocumentError, match="checksum:kb-1"):
        indexer.index_document("kb-1", source, document_id="doc-duplicate")
    second = indexer.index_document("kb-2", source, document_id="doc-2")

    assert first.checksum == second.checksum
    assert len(service.list_documents("kb-1")) == 1
    assert len(service.list_documents("kb-2")) == 1
    assert vector_index.count() > count


def test_same_normalized_text_with_different_raw_bytes_has_different_checksum(tmp_path):
    plain = tmp_path / "plain.txt"
    bom = tmp_path / "bom.txt"
    plain.write_bytes("same text".encode("utf-8"))
    bom.write_bytes("\ufeffsame text".encode("utf-8"))
    indexer, service, _ = make_pipeline()

    first = indexer.index_document("kb-1", plain, document_id="plain")
    second = indexer.index_document("kb-1", bom, document_id="bom")

    assert first.text == second.text == "same text"
    assert first.checksum != second.checksum
    assert len(service.list_documents("kb-1")) == 2


@pytest.mark.parametrize("embedder", [WrongCountEmbedder(), WrongDimensionEmbedder()])
def test_embedding_validation_failure_rolls_back_chunks_vectors_and_marks_failed(
    tmp_path, embedder
):
    source = tmp_path / "failure.txt"
    source.write_text("content to embed", encoding="utf-8")
    vector_index = InMemoryVectorIndex(embedder)
    indexer, service, _ = make_pipeline(
        embedder=embedder,
        vector_index=vector_index,
    )

    with pytest.raises(EmbeddingValidationError):
        indexer.index_document("kb-1", source, document_id="doc-fail")

    document = service.get_document("doc-fail")
    assert document.status == DocumentStatus.FAILED
    assert document.error_message
    assert service.get_document_chunks("doc-fail") == []
    assert vector_index.count() == 0
    assert service.get_knowledge_base("kb-1").status == KnowledgeBaseStatus.FAILED


def test_partial_vector_write_failure_is_compensated(tmp_path):
    source = tmp_path / "vector-failure.txt"
    source.write_text("vector write should fail", encoding="utf-8")
    embedder = HashEmbedder(dim=16)
    vector_index = FailingVectorIndex(embedder)
    indexer, service, _ = make_pipeline(
        embedder=embedder,
        vector_index=vector_index,
    )

    with pytest.raises(IndexingOperationError, match="vector write"):
        indexer.index_document("kb-1", source, document_id="doc-vector-fail")

    assert service.get_document("doc-vector-fail").status == DocumentStatus.FAILED
    assert service.get_document_chunks("doc-vector-fail") == []
    assert vector_index.count() == 0


def test_whitespace_parse_result_rolls_back_and_marks_failed(tmp_path):
    source = tmp_path / "empty.txt"
    source.write_text(" \t\r\n\r\n", encoding="utf-8")
    indexer, service, vector_index = make_pipeline()

    with pytest.raises(EmptyContentError):
        indexer.index_document("kb-1", source, document_id="doc-empty")

    assert service.get_document("doc-empty").status == DocumentStatus.FAILED
    assert service.get_document_chunks("doc-empty") == []
    assert vector_index.count() == 0


def test_empty_file_and_unsupported_format_fail_before_registration(tmp_path):
    empty = tmp_path / "empty.txt"
    unsupported = tmp_path / "data.csv"
    empty.write_bytes(b"")
    unsupported.write_text("value", encoding="utf-8")
    indexer, service, _ = make_pipeline()

    with pytest.raises(EmptyContentError):
        indexer.index_document("kb-1", empty)
    with pytest.raises(UnsupportedDocumentTypeError, match=".csv"):
        indexer.index_document("kb-1", unsupported)

    assert service.list_documents("kb-1") == []
    assert service.get_knowledge_base("kb-1").status == KnowledgeBaseStatus.CREATED


@pytest.mark.parametrize("suffix", [".pdf", ".docx"])
def test_corrupt_structured_document_is_failed_and_cleaned(tmp_path, suffix):
    source = tmp_path / f"broken{suffix}"
    source.write_bytes(b"not a valid structured document")
    indexer, service, vector_index = make_pipeline()

    with pytest.raises(DocumentLoadError):
        indexer.index_document("kb-1", source, document_id="doc-broken")

    assert service.get_document("doc-broken").status == DocumentStatus.FAILED
    assert service.get_document_chunks("doc-broken") == []
    assert vector_index.count() == 0


def test_reindex_replaces_old_chunks_and_vectors_without_duplicates(tmp_path):
    source = tmp_path / "reindex.txt"
    source.write_text("old indexed content", encoding="utf-8")
    indexer, service, vector_index = make_pipeline()
    original = indexer.index_document("kb-1", source, document_id="doc-reindex")
    old_ids = [chunk.id for chunk in service.get_document_chunks(original.id)]

    source.write_text("new replacement content", encoding="utf-8")
    rebuilt = indexer.reindex_document("kb-1", original.id, source)
    first_count = vector_index.count()
    rebuilt_again = indexer.reindex_document("kb-1", original.id, source)
    chunks = service.get_document_chunks(original.id)

    assert rebuilt.id == rebuilt_again.id == original.id
    assert rebuilt.text == "new replacement content"
    assert [chunk.id for chunk in chunks] == old_ids
    assert all("old indexed content" not in chunk.text for chunk in chunks)
    assert len({chunk.id for chunk in chunks}) == len(chunks)
    assert vector_index.count() == first_count == len(chunks)


def test_failed_reindex_removes_old_index_and_marks_rebuilt_document_failed(tmp_path):
    source = tmp_path / "failed-reindex.txt"
    source.write_text("old content", encoding="utf-8")
    indexer, service, vector_index = make_pipeline()
    original = indexer.index_document("kb-1", source, document_id="doc-reindex-fail")
    source.write_text("new content", encoding="utf-8")
    indexer.embedder = WrongCountEmbedder()

    with pytest.raises(EmbeddingValidationError):
        indexer.reindex_document("kb-1", original.id, source)

    rebuilt = service.get_document(original.id)
    assert rebuilt.status == DocumentStatus.FAILED
    assert rebuilt.text == "new content"
    assert service.get_document_chunks(original.id) == []
    assert vector_index.count() == 0


def test_reindex_and_delete_are_isolated_to_target_document(tmp_path):
    first_path = tmp_path / "first.txt"
    second_path = tmp_path / "second.txt"
    first_path.write_text("first original", encoding="utf-8")
    second_path.write_text("second must remain", encoding="utf-8")
    indexer, service, vector_index = make_pipeline()
    first = indexer.index_document("kb-1", first_path, document_id="first")
    second = indexer.index_document("kb-1", second_path, document_id="second")
    second_chunks = service.get_document_chunks(second.id)

    first_path.write_text("first rebuilt", encoding="utf-8")
    indexer.reindex_document("kb-1", first.id, first_path)

    assert service.get_document(second.id).text == "second must remain"
    assert service.get_document_chunks(second.id) == second_chunks
    assert vector_index.count() == (
        len(service.get_document_chunks(first.id)) + len(second_chunks)
    )


def test_missing_file_has_clear_error(tmp_path):
    indexer, _, _ = make_pipeline()

    with pytest.raises(FileNotFoundError, match="was not found"):
        indexer.index_document("kb-1", tmp_path / "missing.txt")
